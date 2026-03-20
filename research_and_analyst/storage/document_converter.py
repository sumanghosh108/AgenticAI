"""
Document Converter — converts Markdown reports to PDF and DOCX formats.

Generates download-ready documents from the report content stored in Supabase.
File names include username for unique identification.
"""

import io
import re
import time
from typing import Optional, Tuple

from research_and_analyst.logger import GLOBAL_LOGGER as log


def _sanitize_filename(text: str) -> str:
    """Create a safe filename from text."""
    safe = re.sub(r'[^\w\s-]', '', text.lower().strip())
    safe = re.sub(r'[\s]+', '_', safe)
    return safe[:60]


def generate_file_name(username: str, topic: str, file_type: str) -> str:
    """
    Generate a unique file name that identifies the user.

    Format: {username}_{topic_slug}_{timestamp}.{ext}
    """
    topic_slug = _sanitize_filename(topic)
    ts = int(time.time())
    return f"{username}_{topic_slug}_{ts}.{file_type}"


def _strip_inline_md(text: str) -> str:
    """Remove inline markdown formatting: **bold**, *italic*, `code`, [link](url)."""
    # Links: [text](url) → text
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
    # Bold+italic: ***text*** or ___text___
    text = re.sub(r'\*{3}(.+?)\*{3}', r'\1', text)
    # Bold: **text** or __text__
    text = re.sub(r'\*{2}(.+?)\*{2}', r'\1', text)
    text = re.sub(r'_{2}(.+?)_{2}', r'\1', text)
    # Italic: *text* or _text_
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    # Inline code: `text`
    text = re.sub(r'`(.+?)`', r'\1', text)
    return text


def markdown_to_pdf(markdown_content: str, title: str = "Analysis Report") -> bytes:
    """
    Convert Markdown content to a PDF byte stream.

    Uses fpdf2 with proper markdown parsing, table rendering,
    and explicit cursor positioning to prevent layout drift.
    """
    from fpdf import FPDF

    pdf = FPDF(orientation="P", unit="mm", format="A4")
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    # Usable width (A4 portrait = 210mm, default margins 10mm each side)
    usable_w = pdf.w - pdf.l_margin - pdf.r_margin

    def _safe(text: str) -> str:
        """Encode to latin-1, replacing bullets and special chars."""
        # Replace unicode bullets with ascii
        text = text.replace("\u2022", "-")
        text = text.replace("\u2019", "'").replace("\u2018", "'")
        text = text.replace("\u201c", '"').replace("\u201d", '"')
        text = text.replace("\u2013", "-").replace("\u2014", "--")
        text = text.replace("\u2265", ">=").replace("\u2264", "<=")
        return text.encode("latin-1", errors="replace").decode("latin-1")

    def _reset_x():
        """Force cursor back to left margin — prevents right-drift."""
        pdf.set_x(pdf.l_margin)

    def _write_text(text: str, font_style: str = "", font_size: int = 10, line_h: float = 6):
        """Write a block of text at the left margin, safely."""
        _reset_x()
        pdf.set_font("Helvetica", font_style, font_size)
        clean = _safe(_strip_inline_md(text))
        if clean.strip():
            pdf.multi_cell(usable_w, line_h, clean)
        _reset_x()

    # ── Title ──
    _reset_x()
    pdf.set_font("Helvetica", "B", 18)
    pdf.multi_cell(usable_w, 12, _safe(title), align="C")
    _reset_x()
    pdf.set_font("Helvetica", "", 9)
    pdf.multi_cell(usable_w, 8, f"Generated: {time.strftime('%Y-%m-%d %H:%M UTC')}", align="C")
    _reset_x()
    pdf.ln(6)

    # ── Table rendering ──
    table_rows: list[list[str]] = []

    def _flush_table():
        nonlocal table_rows
        if not table_rows:
            return

        num_cols = max(len(r) for r in table_rows)
        if num_cols == 0:
            table_rows = []
            return

        # Calculate column widths proportionally based on content
        col_w = usable_w / num_cols
        # Ensure minimum readable column width (at least 15mm)
        if col_w < 15 and num_cols > 3:
            # Switch to smaller font for wide tables
            font_size = max(6, int(8 - (num_cols - 3) * 0.5))
        else:
            font_size = 8
        line_h = font_size * 0.6

        for row_idx, row in enumerate(table_rows):
            y_start = pdf.get_y()

            # Check if we need a new page
            if y_start + 8 > pdf.h - pdf.b_margin:
                pdf.add_page()
                y_start = pdf.get_y()

            # First pass: calculate max row height
            max_cell_h = line_h
            for col_idx in range(num_cols):
                cell_text = _safe(_strip_inline_md(row[col_idx])) if col_idx < len(row) else ""
                if row_idx == 0:
                    pdf.set_font("Helvetica", "B", font_size)
                else:
                    pdf.set_font("Helvetica", "", font_size)
                text_w = pdf.get_string_width(cell_text)
                inner_w = max(col_w - 2, 5)
                n_lines = max(1, int(text_w / inner_w) + 1)
                cell_h = n_lines * line_h + 1
                max_cell_h = max(max_cell_h, cell_h)

            # Second pass: render cells
            for col_idx in range(num_cols):
                cell_text = _safe(_strip_inline_md(row[col_idx])) if col_idx < len(row) else ""
                x_pos = pdf.l_margin + col_idx * col_w
                pdf.set_xy(x_pos, y_start)

                if row_idx == 0:
                    pdf.set_font("Helvetica", "B", font_size)
                else:
                    pdf.set_font("Helvetica", "", font_size)

                # Draw cell border manually, then write text
                pdf.rect(x_pos, y_start, col_w, max_cell_h)
                pdf.set_xy(x_pos + 1, y_start + 0.5)
                pdf.multi_cell(col_w - 2, line_h, cell_text)

            # Move cursor below the row
            pdf.set_xy(pdf.l_margin, y_start + max_cell_h)

        pdf.ln(3)
        _reset_x()
        table_rows = []

    # ── Parse markdown line by line ──
    lines = markdown_content.split("\n")

    for line in lines:
        stripped = line.strip()

        # Flush table if we leave table context
        if table_rows and not stripped.startswith("|"):
            _flush_table()

        # Horizontal rule
        if re.match(r'^-{3,}$', stripped) or re.match(r'^\*{3,}$', stripped):
            _reset_x()
            y = pdf.get_y() + 2
            pdf.line(pdf.l_margin, y, pdf.l_margin + usable_w, y)
            pdf.set_y(y + 3)
            continue

        # H1 — skip (already used as title)
        if stripped.startswith("# ") and not stripped.startswith("## "):
            continue

        # H2
        if stripped.startswith("## "):
            pdf.ln(4)
            heading = _strip_inline_md(stripped.lstrip("#").strip())
            _write_text(heading, "B", 14, 10)
            continue

        # H3
        if stripped.startswith("### ") and not stripped.startswith("#### "):
            pdf.ln(2)
            heading = _strip_inline_md(stripped.lstrip("#").strip())
            _write_text(heading, "B", 12, 8)
            continue

        # H4+
        if stripped.startswith("####"):
            pdf.ln(2)
            heading = _strip_inline_md(stripped.lstrip("#").strip())
            _write_text(heading, "B", 11, 7)
            continue

        # Table separator (skip)
        if re.match(r'^\|[\s\-:|]+\|$', stripped):
            continue

        # Table row
        if stripped.startswith("|") and stripped.endswith("|"):
            cells = [c.strip() for c in stripped.split("|")[1:-1]]
            if cells:
                table_rows.append(cells)
            continue

        # Bullet list
        if stripped.startswith("- ") or stripped.startswith("* "):
            bullet_text = "  - " + _strip_inline_md(stripped[2:])
            _write_text(bullet_text, "", 10, 6)
            continue

        # Numbered list
        m = re.match(r'^(\d+)\.\s+(.+)', stripped)
        if m:
            list_text = f"  {m.group(1)}. " + _strip_inline_md(m.group(2))
            _write_text(list_text, "", 10, 6)
            continue

        # Empty line
        if stripped == "":
            pdf.ln(2)
            continue

        # Body text (strip inline markdown)
        _write_text(stripped, "", 10, 6)

    # Flush any remaining table
    _flush_table()

    return pdf.output()


def _add_rich_paragraph(doc, text: str, base_size: int = 10):
    """Add a paragraph with inline bold/italic rendering to a DOCX document."""
    from docx.shared import Pt

    p = doc.add_paragraph()
    # Split on bold (**...**) and italic (*...*) markers
    parts = re.split(r'(\*{2,3}.+?\*{2,3}|\*[^*]+?\*|`[^`]+?`)', text)

    for part in parts:
        if not part:
            continue
        if part.startswith("***") and part.endswith("***"):
            run = p.add_run(part[3:-3])
            run.bold = True
            run.italic = True
            run.font.size = Pt(base_size)
        elif part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
            run.font.size = Pt(base_size)
        elif part.startswith("*") and part.endswith("*"):
            run = p.add_run(part[1:-1])
            run.italic = True
            run.font.size = Pt(base_size)
        elif part.startswith("`") and part.endswith("`"):
            run = p.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(base_size - 1)
        else:
            # Strip link markdown: [text](url) → text
            cleaned = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', part)
            run = p.add_run(cleaned)
            run.font.size = Pt(base_size)
    return p


def markdown_to_docx(markdown_content: str, title: str = "Analysis Report") -> bytes:
    """
    Convert Markdown content to a DOCX byte stream.

    Uses python-docx for Word document generation with inline formatting.
    """
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Date
    date_para = doc.add_paragraph(f"Generated: {time.strftime('%Y-%m-%d %H:%M UTC')}")
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.runs[0].font.size = Pt(9)

    doc.add_paragraph("")  # Spacer

    # Parse markdown
    lines = markdown_content.split("\n")
    in_table = False
    table_rows = []

    for line in lines:
        stripped = line.strip()

        # Flush table if we exit table context
        if in_table and not stripped.startswith("|"):
            if table_rows:
                _add_table_to_docx(doc, table_rows)
                table_rows = []
            in_table = False

        # Horizontal rule
        if re.match(r'^-{3,}$', stripped) or re.match(r'^\*{3,}$', stripped):
            p = doc.add_paragraph()
            p.add_run("_" * 60)
            continue

        # H1 — skip (already title)
        if stripped.startswith("# ") and not stripped.startswith("## "):
            continue

        # H2
        if stripped.startswith("## ") and not stripped.startswith("### "):
            doc.add_heading(_strip_inline_md(stripped[3:].strip()), level=1)
            continue

        # H3
        if stripped.startswith("### ") and not stripped.startswith("#### "):
            doc.add_heading(_strip_inline_md(stripped[4:].strip()), level=2)
            continue

        # H4+
        if stripped.startswith("####"):
            doc.add_heading(_strip_inline_md(stripped.lstrip("#").strip()), level=3)
            continue

        # Table separator
        if re.match(r'^\|[\s\-:|]+\|$', stripped):
            continue

        # Table row
        if stripped.startswith("|") and stripped.endswith("|"):
            in_table = True
            cells = [c.strip() for c in stripped.split("|")[1:-1]]
            if cells:
                table_rows.append(cells)
            continue

        # Bullet list
        if stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            # Use rich formatting for bullet content
            content = stripped[2:]
            parts = re.split(r'(\*{2}.+?\*{2}|\*[^*]+?\*)', content)
            for part in parts:
                if not part:
                    continue
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith("*") and part.endswith("*"):
                    run = p.add_run(part[1:-1])
                    run.italic = True
                else:
                    p.add_run(re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', part))
            continue

        # Numbered list
        m = re.match(r'^(\d+)\.\s+(.+)', stripped)
        if m:
            p = doc.add_paragraph(style="List Number")
            content = m.group(2)
            parts = re.split(r'(\*{2}.+?\*{2}|\*[^*]+?\*)', content)
            for part in parts:
                if not part:
                    continue
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith("*") and part.endswith("*"):
                    run = p.add_run(part[1:-1])
                    run.italic = True
                else:
                    p.add_run(re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', part))
            continue

        # Empty line
        if stripped == "":
            continue

        # Body text with inline formatting
        _add_rich_paragraph(doc, stripped)

    # Flush any remaining table
    if table_rows:
        _add_table_to_docx(doc, table_rows)

    # Return as bytes
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _add_table_to_docx(doc, rows):
    """Add a table to the DOCX document."""
    if not rows:
        return

    from docx.shared import Pt

    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols, style="Light Grid Accent 1")

    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                cell = table.rows[i].cells[j]
                cell.text = cell_text
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
                # Bold header row
                if i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
