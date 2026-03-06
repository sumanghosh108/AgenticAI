"""
Autonomous Report Generator — LangGraph Workflow

Orchestrates the full research-report pipeline:
  create_analyst → human_feedback → conduct_interview (parallel)
  → write_report / write_introduction / write_conclusion → finalize_report
"""

import os
import json
import re
from datetime import datetime

from langchain_core.messages import HumanMessage
from langchain_community.tools.tavily_search import TavilySearchResults
from langgraph.graph import StateGraph, START, END
from langgraph.types import Send
from langgraph.checkpoint.memory import MemorySaver

from research_and_analyst.schemas import (
    Analyst,
    Perspectives,
    InterviewState,
    ResearchGraphState,
)
from research_and_analyst.prompt_library import (
    ANALYST_CREATION_PROMPT,
    REPORT_WRITER_PROMPT,
    INTRODUCTION_PROMPT,
    CONCLUSION_PROMPT,
)
from research_and_analyst.workflows.interview_workflow import InterviewGraphBuilder
from research_and_analyst.logger import GLOBAL_LOGGER as log
from research_and_analyst.exception.custom_exception import ResearchAnalystException


class AutonomousReportGenerator:
    """End-to-end autonomous report generator powered by LangGraph."""

    def __init__(self, llm):
        """
        Args:
            llm: A LangChain chat model instance (ChatOpenAI, ChatGoogleGenerativeAI, etc.)
        """
        self.llm = llm
        self.tavily_search = TavilySearchResults(max_results=3)
        self.logger = log.bind(module="AutonomousReportGenerator")
        self.memory = MemorySaver()

    # ─── Node: create analyst personas ────────────────────────────────

    def _parse_analysts_from_text(self, text: str, max_analysts: int) -> list[Analyst]:
        """Parse analyst personas from LLM text response (JSON or freeform)."""

        # Try to extract JSON array from response
        json_match = re.search(r'\[.*\]', text, re.DOTALL)
        if json_match:
            try:
                raw = json.loads(json_match.group())
                analysts = []
                for item in raw[:max_analysts]:
                    analysts.append(Analyst(
                        name=item.get("name", "Unknown"),
                        role=item.get("role", "Analyst"),
                        affiliation=item.get("affiliation", "Independent"),
                        description=item.get("description", "Research analyst"),
                    ))
                return analysts
            except (json.JSONDecodeError, KeyError):
                pass

        # Fallback: generate default analysts
        self.logger.warning("Could not parse JSON from LLM; using fallback analysts")
        return [
            Analyst(
                name=f"Analyst {i+1}",
                role="Research Analyst",
                affiliation="Independent Research",
                description=f"General research analyst #{i+1} for this topic",
            )
            for i in range(max_analysts)
        ]

    def create_analyst(self, state: ResearchGraphState):
        """Generate analyst personas using the LLM."""
        try:
            topic = state.get("topic", "General Research")
            max_analysts = state.get("max_analysts", 3)

            self.logger.info(
                "Creating analyst personas",
                topic=topic,
                max_analysts=max_analysts,
            )

            prompt = ANALYST_CREATION_PROMPT.format(
                topic=topic,
                max_analysts=max_analysts,
            )

            # Add JSON instruction (works with any model, no response_format needed)
            json_instruction = (
                "\n\nReturn your answer as a JSON array of objects, each with keys: "
                '"name", "role", "affiliation", "description". '
                "Return ONLY the JSON array, no other text."
            )

            response = self.llm.invoke([HumanMessage(content=prompt + json_instruction)])
            analysts = self._parse_analysts_from_text(response.content, max_analysts)

            for a in analysts:
                self.logger.info("Analyst created", name=a.name, role=a.role)

            return {"analysts": analysts}

        except Exception as e:
            self.logger.error("Error creating analysts", error=str(e))
            raise ResearchAnalystException("Failed to create analyst personas", e)

    # ─── Node: human feedback gate ────────────────────────────────────

    def human_feedback(self, state: ResearchGraphState):
        """Pass-through node; the graph interrupts here for human review."""
        pass

    # ─── Node: write the report body ──────────────────────────────────

    def write_report(self, state: ResearchGraphState):
        """Compile all interview sections into a cohesive report body."""
        try:
            topic = state.get("topic", "Untitled Topic")
            sections = state.get("sections", [])

            self.logger.info("Writing report body", num_sections=len(sections))

            sections_text = "\n\n---\n\n".join(sections) if sections else "No sections available."

            prompt = REPORT_WRITER_PROMPT.format(
                topic=topic,
                sections=sections_text,
            )

            response = self.llm.invoke([HumanMessage(content=prompt)])

            self.logger.info("Report body written successfully")
            return {"content": response.content}

        except Exception as e:
            self.logger.error("Error writing report", error=str(e))
            raise ResearchAnalystException("Failed to write report body", e)

    # ─── Node: write the introduction ─────────────────────────────────

    def write_introduction(self, state: ResearchGraphState):
        """Generate the report introduction."""
        try:
            topic = state.get("topic", "Untitled Topic")
            sections = state.get("sections", [])

            self.logger.info("Writing introduction")

            sections_text = "\n\n".join(sections) if sections else "No sections available."

            prompt = INTRODUCTION_PROMPT.format(
                topic=topic,
                sections=sections_text,
            )

            response = self.llm.invoke([HumanMessage(content=prompt)])

            self.logger.info("Introduction written successfully")
            return {"introduction": response.content}

        except Exception as e:
            self.logger.error("Error writing introduction", error=str(e))
            raise ResearchAnalystException("Failed to write introduction", e)

    # ─── Node: write the conclusion ───────────────────────────────────

    def write_conclusion(self, state: ResearchGraphState):
        """Generate the report conclusion."""
        try:
            topic = state.get("topic", "Untitled Topic")
            sections = state.get("sections", [])

            self.logger.info("Writing conclusion")

            # Use sections as the report context for the conclusion
            report_text = "\n\n".join(sections) if sections else "No content available."

            prompt = CONCLUSION_PROMPT.format(
                topic=topic,
                report=report_text,
            )

            response = self.llm.invoke([HumanMessage(content=prompt)])

            self.logger.info("Conclusion written successfully")
            return {"conclusion": response.content}

        except Exception as e:
            self.logger.error("Error writing conclusion", error=str(e))
            raise ResearchAnalystException("Failed to write conclusion", e)

    # ─── Node: assemble final report ──────────────────────────────────

    def finalize_report(self, state: ResearchGraphState):
        """Concatenate introduction + body + conclusion into the final report."""
        try:
            introduction = state.get("introduction", "")
            content = state.get("content", "")
            conclusion = state.get("conclusion", "")
            topic = state.get("topic", "Untitled Topic")

            self.logger.info("Finalising report")

            final = (
                f"# Research Report: {topic}\n\n"
                f"{introduction}\n\n"
                f"---\n\n"
                f"{content}\n\n"
                f"---\n\n"
                f"{conclusion}\n"
            )

            self.logger.info("Report finalised", length=len(final))
            return {"final_report": final}

        except Exception as e:
            self.logger.error("Error finalising report", error=str(e))
            raise ResearchAnalystException("Failed to finalise report", e)

    # ─── Utility: save the report to disk ─────────────────────────────

    def save_report(self, content: str, topic: str, fmt: str = "md") -> str:
        """
        Save the final report to the generated_report/ directory.

        Args:
            content: The full report text.
            topic: Topic name (used for folder naming).
            fmt: File format — 'md', 'docx', or 'pdf'.

        Returns:
            str: Absolute path to the saved file.
        """
        try:
            # Sanitise topic for use as a directory name
            safe_topic = "".join(
                c if c.isalnum() or c in (" ", "-", "_") else "_"
                for c in topic
            ).strip().replace(" ", "_")

            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            report_dir = os.path.join(
                os.getcwd(), "generated_report", f"{safe_topic}_{timestamp}"
            )
            os.makedirs(report_dir, exist_ok=True)

            file_name = f"report.{fmt}"
            file_path = os.path.join(report_dir, file_name)

            if fmt == "md":
                with open(file_path, "w", encoding="utf-8") as f:
                    f.write(content)

            elif fmt == "docx":
                try:
                    from docx import Document

                    doc = Document()
                    for line in content.split("\n"):
                        if line.startswith("# "):
                            doc.add_heading(line[2:], level=1)
                        elif line.startswith("## "):
                            doc.add_heading(line[3:], level=2)
                        elif line.startswith("### "):
                            doc.add_heading(line[4:], level=3)
                        elif line.strip() == "---":
                            doc.add_page_break()
                        elif line.strip():
                            doc.add_paragraph(line)
                    doc.save(file_path)
                except ImportError:
                    self.logger.warning(
                        "python-docx not installed; falling back to .md"
                    )
                    file_path = os.path.join(report_dir, "report.md")
                    with open(file_path, "w", encoding="utf-8") as f:
                        f.write(content)

            elif fmt == "pdf":
                try:
                    from fpdf import FPDF

                    # Sanitise Unicode chars that Helvetica can't render
                    def _sanitize(text: str) -> str:
                        replacements = {
                            "\u2014": "--",   # em dash
                            "\u2013": "-",    # en dash
                            "\u2018": "'",    # left single quote
                            "\u2019": "'",    # right single quote
                            "\u201c": '"',    # left double quote
                            "\u201d": '"',    # right double quote
                            "\u2026": "...",  # ellipsis
                            "\u2022": "*",    # bullet
                            "\u00a0": " ",    # non-breaking space
                        }
                        for uni, asc in replacements.items():
                            text = text.replace(uni, asc)
                        # Strip any remaining non-latin-1 chars
                        return text.encode("latin-1", errors="replace").decode("latin-1")

                    pdf = FPDF()
                    pdf.set_auto_page_break(auto=True, margin=15)
                    pdf.add_page()

                    for line in content.split("\n"):
                        stripped = _sanitize(line.strip())

                        if stripped.startswith("# "):
                            pdf.set_font("Helvetica", "B", 18)
                            pdf.cell(0, 12, stripped[2:], new_x="LMARGIN", new_y="NEXT")
                            pdf.ln(4)
                        elif stripped.startswith("## "):
                            pdf.set_font("Helvetica", "B", 15)
                            pdf.cell(0, 10, stripped[3:], new_x="LMARGIN", new_y="NEXT")
                            pdf.ln(3)
                        elif stripped.startswith("### "):
                            pdf.set_font("Helvetica", "B", 13)
                            pdf.cell(0, 9, stripped[4:], new_x="LMARGIN", new_y="NEXT")
                            pdf.ln(2)
                        elif stripped == "---":
                            pdf.ln(6)
                            pdf.set_draw_color(180, 180, 180)
                            pdf.line(10, pdf.get_y(), 200, pdf.get_y())
                            pdf.ln(6)
                        elif stripped:
                            pdf.set_font("Helvetica", "", 11)
                            pdf.multi_cell(0, 6, stripped)
                            pdf.ln(2)

                    pdf.output(file_path)
                except ImportError:
                    self.logger.warning(
                        "fpdf2 not installed; falling back to .md"
                    )
                    file_path = os.path.join(report_dir, "report.md")
                    with open(file_path, "w", encoding="utf-8") as f:
                        f.write(content)

            self.logger.info("Report saved", path=file_path, format=fmt)
            return file_path

        except Exception as e:
            self.logger.error("Error saving report", error=str(e))
            raise ResearchAnalystException("Failed to save report", e)

    # ─── Build the full report-generation graph ───────────────────────

    def build_graph(self):
        """
        Build and compile the outer report-generation StateGraph.

        Returns:
            CompiledGraph: The compiled LangGraph workflow.
        """
        try:
            self.logger.info("Building report generation graph")

            builder = StateGraph(ResearchGraphState)
            interview_graph = InterviewGraphBuilder(
                self.llm, self.tavily_search
            ).build()

            # ── Local helper: fan-out to parallel interviews ──────────
            def initiate_all_interviews(state: ResearchGraphState):
                topic = state.get("topic", "Untitled Topic")
                analysts = state.get("analysts", [])
                human_feedback = state.get("human_analyst_feedback", "")

                # If feedback is empty / "approve", proceed to interviews
                # otherwise loop back to regenerate analysts
                if not human_feedback or human_feedback.strip().lower() in (
                    "",
                    "approve",
                    "approved",
                    "ok",
                    "yes",
                ):
                    if not analysts:
                        self.logger.warning(
                            "No analysts found — skipping interviews"
                        )
                        return END

                    return [
                        Send(
                            "conduct_interview",
                            {
                                "analyst": analyst,
                                "messages": [
                                    HumanMessage(
                                        content=f"So, let's discuss about {topic}."
                                    )
                                ],
                                "max_num_turns": 2,
                                "context": [],
                                "interview": "",
                                "sections": [],
                            },
                        )
                        for analyst in analysts
                    ]

                # Feedback not approved → loop back
                return "create_analyst"

            # ── Register nodes ────────────────────────────────────────
            builder.add_node("create_analyst", self.create_analyst)
            builder.add_node("human_feedback", self.human_feedback)
            builder.add_node("conduct_interview", interview_graph)
            builder.add_node("write_report", self.write_report)
            builder.add_node("write_introduction", self.write_introduction)
            builder.add_node("write_conclusion", self.write_conclusion)
            builder.add_node("finalize_report", self.finalize_report)

            # ── Edges ─────────────────────────────────────────────────
            builder.add_edge(START, "create_analyst")
            builder.add_edge("create_analyst", "human_feedback")

            builder.add_conditional_edges(
                "human_feedback",
                initiate_all_interviews,
                ["conduct_interview", "create_analyst", END],
            )

            builder.add_edge("conduct_interview", "write_report")
            builder.add_edge("conduct_interview", "write_introduction")
            builder.add_edge("conduct_interview", "write_conclusion")

            builder.add_edge(
                ["write_report", "write_introduction", "write_conclusion"],
                "finalize_report",
            )

            builder.add_edge("finalize_report", END)

            # ── Compile with interrupt for human feedback ─────────────
            graph = builder.compile(
                interrupt_before=["human_feedback"],
                checkpointer=self.memory,
            )

            self.logger.info("Report generation graph built successfully")
            return graph

        except Exception as e:
            self.logger.error("Error building report graph", error=str(e))
            raise ResearchAnalystException(
                "Failed to build report generation graph", e
            )


# ─────────────────────────────────────────────────────────────────────
# Standalone test
# ─────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    import os
    from research_and_analyst.utils.model_loader import ModelLoader

    os.environ.setdefault("LLM_PROVIDER", "openrouter")

    loader = ModelLoader()
    llm = loader.load_llm()

    generator = AutonomousReportGenerator(llm)
    graph = generator.build_graph()

    print("✅ Report generation graph compiled successfully")
    print(f"   Nodes: {list(graph.get_graph().nodes.keys())}")
